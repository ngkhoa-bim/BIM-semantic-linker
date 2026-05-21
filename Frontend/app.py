# =============================================================
# BIM Semantic Linker — Frontend Streamlit v2.0
# =============================================================
# Kiến trúc: Frontend KHÔNG ghi trực tiếp vào Neo4j.
# Mọi thao tác ghi đều đi qua Backend API.
# =============================================================

import os
import uuid
import io
import zipfile

import requests
import pandas as pd
import PyPDF2
import docx
import streamlit as st
from itertools import groupby

# ──────────────────────────────────────────────────────────────
# Cấu hình trang
# ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="BIM Semantic Linker",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ──────────────────────────────────────────────────────────────
# Backend URL — ưu tiên: st.secrets > env var > localhost
# ──────────────────────────────────────────────────────────────
def get_backend_url() -> str:
    env_url = os.environ.get("BACKEND_URL")
    if env_url:
        return env_url.rstrip("/")

    try:
        return st.secrets.get("BACKEND_URL", "http://localhost:7860").rstrip("/")
    except Exception:
        return "http://localhost:7860"

BACKEND = get_backend_url()

# ──────────────────────────────────────────────────────────────
# Hàm trích xuất văn bản từ các định dạng tài liệu
# ──────────────────────────────────────────────────────────────
def extract_pdf(file) -> str:
    """Trích xuất toàn bộ text từ PDF (mỗi trang nối nhau)."""
    reader = PyPDF2.PdfReader(file)
    pages  = []
    for i, page in enumerate(reader.pages):
        txt = page.extract_text()
        if txt and txt.strip():
            pages.append(f"--- Trang {i + 1} ---\n{txt}")
    return "\n\n".join(pages)


def extract_docx(file) -> str:
    """
    Trích xuất text từ DOCX: đoạn văn + nội dung bảng.
    """
    doc   = docx.Document(file)
    parts = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_xlsx(file) -> str:
    """
    Trích xuất tất cả sheets từ XLSX thành Markdown.
    Markdown giúp Claude dễ đọc dữ liệu bảng hơn.
    """
    xls   = pd.ExcelFile(file)
    parts = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name).dropna(how="all")
        if df.empty:
            continue
        parts.append(f"=== Sheet: {sheet_name} ===")
        try:
            parts.append(df.to_markdown(index=False))
        except Exception:
            parts.append(df.to_string(index=False))
    return "\n\n".join(parts)


def extract_csv(file) -> str:
    """Trích xuất CSV thành Markdown."""
    df = pd.read_csv(file).dropna(how="all")
    try:
        return df.to_markdown(index=False)
    except Exception:
        return df.to_string(index=False)


def extract_text(uploaded_file) -> str:
    """Dispatcher: chọn hàm trích xuất theo extension."""
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_pdf(io.BytesIO(uploaded_file.read()))
    elif name.endswith(".docx"):
        return extract_docx(io.BytesIO(uploaded_file.read()))
    elif name.endswith(".xlsx"):
        return extract_xlsx(io.BytesIO(uploaded_file.read()))
    elif name.endswith(".csv"):
        return extract_csv(io.BytesIO(uploaded_file.read()))
    return ""



# ──────────────────────────────────────────────────────────────
# ICDD ZIP helpers — validate và chuẩn hoá package người dùng upload
# ──────────────────────────────────────────────────────────────
def _is_meaningful_zip_file(name: str) -> bool:
    """Bỏ qua folder entry, file hệ thống và placeholder rỗng."""
    clean = name.replace("\\", "/").strip("/")
    if not clean or clean.endswith("/"):
        return False
    parts = clean.split("/")
    if any(part in {"__MACOSX", ".DS_Store", "Thumbs.db"} for part in parts):
        return False
    if parts[-1] in {".gitkeep", ".DS_Store", "Thumbs.db"}:
        return False
    return True


def _strip_single_root_folder(names: list[str]) -> tuple[list[str], str | None]:
    """
    Nếu người dùng nén cả thư mục cha, ví dụ:
        SNT_ICDD_starter/Container.rdf
    thì tự động chuẩn hoá về:
        Container.rdf
    để file .icdd cuối cùng đúng cấu trúc root của ICDD.
    """
    normalized = [n.replace("\\", "/").strip("/") for n in names]
    if "Container.rdf" in normalized:
        return normalized, None

    first_parts = [n.split("/", 1)[0] for n in normalized if "/" in n]
    if not first_parts:
        return normalized, None

    unique_roots = set(first_parts)
    if len(unique_roots) == 1:
        root = next(iter(unique_roots))
        candidate = f"{root}/Container.rdf"
        if candidate in normalized:
            stripped = [n[len(root) + 1:] if n.startswith(root + "/") else n for n in normalized]
            return stripped, root

    return normalized, None


def validate_and_normalize_icdd_zip(zip_bytes: bytes, project_id: str) -> dict:
    """
    Kiểm tra ZIP đã bổ sung payload/ontology và trả về bytes đã chuẩn hoá.
    Không giải nén ra ổ đĩa; chỉ xử lý trong bộ nhớ để an toàn và phù hợp Render.
    """
    try:
        src = zipfile.ZipFile(io.BytesIO(zip_bytes), "r")
    except zipfile.BadZipFile:
        return {
            "ok": False,
            "errors": ["File upload không phải ZIP hợp lệ."],
            "warnings": [],
            "summary": {},
            "icdd_bytes": None,
        }

    infos = src.infolist()
    meaningful_raw = []
    for info in infos:
        original_name = info.filename.replace("\\", "/")
        if info.is_dir():
            continue
        if _is_meaningful_zip_file(original_name):
            meaningful_raw.append(original_name.strip("/"))
    normalized_names, stripped_root = _strip_single_root_folder(meaningful_raw)

    # Map tên đã chuẩn hoá -> tên gốc trong ZIP để copy bytes sang package cuối.
    name_map = {}
    if stripped_root:
        for raw in meaningful_raw:
            if raw.startswith(stripped_root + "/"):
                name_map[raw[len(stripped_root) + 1:]] = raw
            else:
                name_map[raw] = raw
    else:
        name_map = {raw: raw for raw in meaningful_raw}

    names = sorted(name_map.keys())
    errors = []
    warnings = []

    has_container = "Container.rdf" in names
    linkset_files = [
        n for n in names
        if n.startswith("Linkset Documents/")
        and n.lower().endswith((".ttl", ".rdf", ".jsonld"))
    ]
    payload_files = [
        n for n in names
        if n.startswith("Payload Documents/")
        and _is_meaningful_zip_file(n)
    ]
    ontology_files = [
        n for n in names
        if n.startswith("Ontology Resources/")
        and _is_meaningful_zip_file(n)
    ]

    if not has_container:
        errors.append("Thiếu `Container.rdf` ở root của ZIP.")
    if not linkset_files:
        errors.append("Thiếu file linkset trong `Linkset Documents/`, ví dụ `linkset_{}.ttl`.".format(project_id))
    if not payload_files:
        errors.append("`Payload Documents/` chưa có tài liệu gốc PDF/DOCX/XLSX/IFC.")
    if not ontology_files:
        errors.append("`Ontology Resources/` chưa có ontology tham chiếu, ví dụ `bot.ttl`, `IFC2X3_Final.ttl`.")

    if stripped_root:
        warnings.append(
            f"Đã phát hiện ZIP có thư mục cha `{stripped_root}/`; app sẽ tự bỏ lớp này khi xuất `.icdd`."
        )

    expected_linkset = f"Linkset Documents/linkset_{project_id}.ttl"
    if linkset_files and expected_linkset not in linkset_files:
        warnings.append(
            f"Không thấy `{expected_linkset}`; vẫn chấp nhận vì có file linkset RDF khác."
        )

    lower_ontology = {n.split("/")[-1].lower() for n in ontology_files}
    if ontology_files and "bot.ttl" not in lower_ontology:
        warnings.append("Khuyến nghị bổ sung `bot.ttl` vào `Ontology Resources/`.")
    if ontology_files and not any("ifc" in f and f.endswith(".ttl") for f in lower_ontology):
        warnings.append("Khuyến nghị bổ sung ontology IFC, ví dụ `IFC2X3_Final.ttl`, vào `Ontology Resources/`.")

    summary = {
        "Container.rdf": 1 if has_container else 0,
        "Linkset Documents": len(linkset_files),
        "Payload Documents": len(payload_files),
        "Ontology Resources": len(ontology_files),
    }

    if errors:
        return {
            "ok": False,
            "errors": errors,
            "warnings": warnings,
            "summary": summary,
            "icdd_bytes": None,
        }

    # Ghi lại ZIP đã chuẩn hoá. File `.icdd` thực chất là ZIP, chỉ khác phần mở rộng.
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        # đảm bảo các thư mục chính luôn tồn tại
        for folder in ["Payload Documents/", "Linkset Documents/", "Ontology Resources/"]:
            dst.writestr(folder, b"")

        for normalized_name in names:
            src_name = name_map[normalized_name]
            dst.writestr(normalized_name, src.read(src_name))

    return {
        "ok": True,
        "errors": [],
        "warnings": warnings,
        "summary": summary,
        "icdd_bytes": out.getvalue(),
        "files": names,
    }

# ──────────────────────────────────────────────────────────────
# Backend API helpers — tất cả request đều qua đây
# ──────────────────────────────────────────────────────────────
def api_health() -> dict:
    try:
        r = requests.get(f"{BACKEND}/health", timeout=8)
        r.raise_for_status()
        return r.json()
    except Exception as e:
        return {"error": str(e)}


def api_ingest_ifc(project_id: str, file_bytes: bytes, filename: str) -> dict:
    try:
        r = requests.post(
            f"{BACKEND}/api/projects/{project_id}/ifc/ingest",
            files={"file": (filename, file_bytes, "application/octet-stream")},
            timeout=600,
        )
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        return {"error": f"HTTP {e.response.status_code}: {e.response.text[:200]}"}
    except Exception as e:
        return {"error": str(e)}


def api_get_anchors(project_id: str) -> dict:
    try:
        r = requests.get(
            f"{BACKEND}/api/projects/{project_id}/anchors",
            timeout=30,
        )
        r.raise_for_status()
        return r.json()
    except Exception as e:
        return {"error": str(e)}

def api_analyze(project_id: str, doc_name: str, doc_text: str) -> dict:
    try:
        r = requests.post(
            f"{BACKEND}/api/projects/{project_id}/documents/analyze",
            json={"document_name": doc_name, "document_text": doc_text},
            timeout=300,
        )
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        return {"error": f"HTTP {e.response.status_code}: {e.response.text[:200]}"}
    except Exception as e:
        return {"error": str(e)}


def api_confirm_link(project_id: str, payload: dict) -> dict:
    try:
        r = requests.post(
            f"{BACKEND}/api/projects/{project_id}/links/confirm",
            json=payload,
            timeout=30,
        )
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        return {"error": f"HTTP {e.response.status_code}: {e.response.text[:200]}"}
    except Exception as e:
        return {"error": str(e)}

def api_export_linkset(project_id: str, fmt: str = "turtle") -> requests.Response | None:
    """
    Gọi GET /api/projects/<pid>/links/export?format=<fmt>
    Trả về requests.Response thô (để lấy .content bytes),
    hoặc None nếu có lỗi kết nối.
    """
    try:
        r = requests.get(
            f"{BACKEND}/api/projects/{project_id}/links/export",
            params={"format": fmt, "save": "false"},
            timeout=120,
        )
        r.raise_for_status()
        return r
    except requests.HTTPError as e:
        st.error(f"❌ HTTP {e.response.status_code}: {e.response.text[:300]}")
        return None
    except Exception as e:
        st.error(f"❌ Lỗi kết nối: {e}")
        return None

def api_export_container(project_id: str) -> requests.Response | None:
    """
    Gọi GET /api/projects/<pid>/container/export
    Trả về requests.Response thô chứa nội dung Container.rdf,
    hoặc None nếu có lỗi.
    """
    try:
        r = requests.get(
            f"{BACKEND}/api/projects/{project_id}/container/export",
            params={"save": "false"},
            timeout=60,
        )
        r.raise_for_status()
        return r
    except requests.HTTPError as e:
        st.error(f"❌ HTTP {e.response.status_code}: {e.response.text[:300]}")
        return None
    except Exception as e:
        st.error(f"❌ Lỗi kết nối: {e}")
        return None

def api_package_icdd(project_id: str) -> requests.Response | None:
    """
    Gọi GET /api/projects/<pid>/icdd/package
    Trả về requests.Response thô chứa file ZIP (.icdd),
    hoặc None nếu có lỗi.
    """
    try:
        r = requests.get(
            f"{BACKEND}/api/projects/{project_id}/icdd/package",
            timeout=180,
        )
        r.raise_for_status()
        return r
    except requests.HTTPError as e:
        st.error(f"❌ HTTP {e.response.status_code}: {e.response.text[:300]}")
        return None
    except Exception as e:
        st.error(f"❌ Lỗi kết nối: {e}")
        return None

# ─────────────────────────────────────────────────────────────────────────────
#   - Thêm tham số doc_id, doc_name, suggestions_data
#   - Backend cần suggestions_data để ghi DocumentLink đầy đủ vào Neo4j
#     (element_global_id, confidence, method, evidence, v.v.)
#   - Khi action="REJECTED", suggestions_data không cần thiết → truyền [] là được
# ─────────────────────────────────────────────────────────────────────────────
def api_confirm_group( project_id: str, suggestion_ids: list, action: str, doc_id: str = "", doc_name: str = "", suggestions_data: list = None,) -> dict:
    """
    Xác nhận hoặc từ chối một nhóm suggestions cùng lúc.
    Args:
        project_id:      ID dự án trong Neo4j
        suggestion_ids:  Danh sách suggestion_id cần xử lý
        action:          "CONFIRMED" hoặc "REJECTED"
        doc_id:          document_id từ session_state (cần khi CONFIRMED)
        doc_name:        Tên file tài liệu (cần khi CONFIRMED)
        suggestions_data: Dữ liệu đầy đủ các suggestion (cần khi CONFIRMED để Backend biết element_global_id, confidence, v.v.)
    """
    try:
        r = requests.post(
            f"{BACKEND}/api/projects/{project_id}/links/confirm-group",
            json={
                "suggestion_ids":  suggestion_ids,
                "action":          action,
                "document_id":     doc_id,
                "document_name":   doc_name,
                "suggestions_data": suggestions_data or [],
            },
            timeout=60,  # group có thể ghi nhiều records → timeout dài hơn single confirm
        )
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        return {"error": f"HTTP {e.response.status_code}: {e.response.text[:200]}"}
    except Exception as e:
        return {"error": str(e)}


def api_get_links(project_id: str) -> dict:
    try:
        r = requests.get(
            f"{BACKEND}/api/projects/{project_id}/links",
            timeout=30,
        )
        r.raise_for_status()
        return r.json()
    except Exception as e:
        return {"error": str(e)}


# ──────────────────────────────────────────────────────────────
# Session state — khởi tạo 1 lần, giữ qua mọi rerun
# ──────────────────────────────────────────────────────────────
_STATE_DEFAULTS = {
    "suggestions":        [],      # Danh sách đề xuất AI từ analyze
    "extracted_text":     "",      # Văn bản đã trích xuất từ tài liệu
    "current_file_name":  None,    # Tên file đang xử lý (để tránh re-extract)
    "current_doc_id":     None,    # document_id ổn định từ Backend
    "confirmed_uids":     set(),   # ui_id của các đề xuất đã được confirm
    "skipped_uids":       set(),   # ui_id của các đề xuất đã bỏ qua
}
for _k, _v in _STATE_DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ──────────────────────────────────────────────────────────────
# Sidebar
# ──────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## ⚙️ Cấu hình")

    project_id = st.text_input(
        "🔑 Project ID",
        value="P01",
        help="Định danh dự án trong Neo4j. Mỗi IFC file nên có project_id riêng.",
    )

    st.divider()
    st.caption("🌐 Backend URL đang dùng")
    st.code(BACKEND, language=None)

    col_ping, col_blank = st.columns(2)
    if col_ping.button("🔍 Ping Backend"):
        with st.spinner("Đang kiểm tra..."):
            res = api_health()
        if "error" in res:
            st.error(f"❌ {res['error']}")
        else:
            st.success(f"✅ {res.get('status','ok')} — v{res.get('version','?')}")

    st.divider()

    # Thống kê phiên làm việc
    st.markdown("### 📊 Phiên làm việc")
    m1, m2, m3 = st.columns(3)
    m1.metric("Đề xuất",  len(st.session_state.suggestions))
    m2.metric("Đã xác nhận", len(st.session_state.confirmed_uids))
    m3.metric("Bỏ qua",   len(st.session_state.skipped_uids))

    if st.button("🔄 Reset phiên", help="Xoá suggestions và bộ đệm văn bản"):
        for k, v in _STATE_DEFAULTS.items():
            st.session_state[k] = v if not isinstance(v, set) else set()
        st.rerun()


# ──────────────────────────────────────────────────────────────
# Main UI
# ──────────────────────────────────────────────────────────────
st.title("🏗️ BIM Semantic Linker")
st.caption(
    "Liên kết ngữ nghĩa giữa mô hình IFC và tài liệu dự án "
    "— ID-based Matching + Semantic Discovery + Human Validation"
)

tab_ifc, tab_docs, tab_links, tab_icdd = st.tabs([
    "🏗️ 1. Nạp Mô hình IFC",
    "📄 2. Phân tích Tài liệu",
    "🔗 3. Liên kết đã xác nhận",
    "📦 4. Xuất ICDD / Linkset",
])


# ══════════════════════════════════════════════════════════════
# TAB 1 — IFC Ingestion
# ══════════════════════════════════════════════════════════════
with tab_ifc:
    st.subheader("📦 Nạp file IFC vào Graph Database")
    st.info(
        "Upload file `.ifc` để Backend (IfcOpenShell) trích xuất "
        "cấu kiện, Psets, spatial info và nạp vào Neo4j Aura.",
        icon="ℹ️",
    )

    ifc_file = st.file_uploader(
        "Chọn file IFC",
        type=["ifc"],
        help="Chỉ hỗ trợ định dạng IFC (IFC2x3, IFC4)",
    )

    if ifc_file:
        st.write(f"📎 **{ifc_file.name}** — {ifc_file.size / 1024:.1f} KB")

        if st.button("🚀 Ingest vào Neo4j", type="primary"):
            with st.spinner(f"Đang xử lý `{ifc_file.name}`… Quá trình này có thể mất 1–3 phút."):
                result = api_ingest_ifc(project_id, ifc_file.getvalue(), ifc_file.name)

            if "error" in result:
                st.error(f"❌ Lỗi: {result['error']}")
            else:
                st.success("✅ Ingest thành công!")
                c1, c2, c3 = st.columns(3)
                c1.metric("Cấu kiện đã nạp", result.get("ingested_count", 0))
                c2.metric("Lỗi bỏ qua",      result.get("error_count", 0))
                c3.metric("Project ID",       result.get("project_id", "-"))

                if result.get("errors"):
                    with st.expander("⚠️ Xem chi tiết lỗi (tối đa 10)"):
                        for err in result["errors"]:
                            st.warning(
                                f"**GlobalId:** `{err.get('globalId','?')}` "
                                f"→ {err.get('error','')}"
                            )

    # Preview anchor list sau khi ingest
    st.divider()
    st.subheader("🔍 Preview Anchor List")
    if st.button("Tải danh sách cấu kiện từ Neo4j"):
        with st.spinner("Đang truy vấn Neo4j..."):
            anchor_res = api_get_anchors(project_id)
        if "error" in anchor_res:
            st.error(anchor_res["error"])
        else:
            anchors = anchor_res.get("anchors", [])
            st.success(f"Tìm thấy **{len(anchors)}** cấu kiện trong project `{project_id}`")
            if anchors:
                df_anchors = pd.DataFrame(anchors)
                # Hiển thị các cột quan trọng
                cols_show = [c for c in ["globalId", "name", "ifcType", "mark", "reference", "tag", "storey"] if c in df_anchors.columns]
                st.dataframe(df_anchors[cols_show], use_container_width=True)

# ══════════════════════════════════════════════════════════════
# TAB 2 — Document Analysis
# ══════════════════════════════════════════════════════════════
with tab_docs:
    st.subheader("📄 Phân tích Tài liệu & Đề xuất Liên kết")

    uploaded = st.file_uploader(
        "Tải tài liệu dự án",
        type=["pdf", "docx", "xlsx", "csv"],
        help="Hỗ trợ: PDF, Word (.docx), Excel (.xlsx), CSV",
    )

    # ── Reset state khi người dùng xoá file ──────────────────
    if uploaded is None:
        if st.session_state.current_file_name is not None:
            st.session_state.current_file_name = None
            st.session_state.extracted_text    = ""
            st.session_state.suggestions       = []

    # ── Trích xuất text khi có file mới ──────────────────────
    elif uploaded.name != st.session_state.current_file_name:
        with st.spinner(f"Đang trích xuất văn bản từ `{uploaded.name}`..."):
            try:
                # Đọc bytes một lần, tránh SeekableStream bị consumed
                text = extract_text(uploaded)

                st.session_state.extracted_text    = text
                st.session_state.current_file_name = uploaded.name
                st.session_state.suggestions       = []

                if not text.strip():
                    st.warning(
                        "⚠️ Không trích xuất được văn bản. "
                        "File có thể là scan/image PDF hoặc bị bảo vệ."
                    )
                else:
                    st.success(f"✅ Trích xuất xong: **{len(text):,}** ký tự")

            except Exception as e:
                st.error(f"❌ Lỗi khi đọc file: {e}")

    # ── Hiển thị preview + nút phân tích ─────────────────────
    if st.session_state.extracted_text:
        with st.expander("👁️ Xem trước nội dung đã trích xuất", expanded=False):
            preview = st.session_state.extracted_text
            if len(preview) > 3000:
                preview = preview[:3000] + "\n\n… [đã rút gọn, hiển thị 3000 ký tự đầu]"
            st.text(preview)

        st.info(
            f"📝 Tài liệu: `{st.session_state.current_file_name}` — "
            f"**{len(st.session_state.extracted_text):,}** ký tự"
        )

        col_btn, col_info = st.columns([2, 5])
        if col_btn.button("🤖 Phân tích với AI", type="primary"):
            with st.spinner(
                "Đang chạy ID-based Matching + Semantic Discovery (Claude)… "
                "Vui lòng chờ 15–60 giây."
            ):
                result = api_analyze(
                    project_id,
                    st.session_state.current_file_name,
                    st.session_state.extracted_text,
                )

            if "error" in result:
                st.error(f"❌ Lỗi phân tích: {result['error']}")
            else:
                doc_info = result.get("document", {})
                st.session_state.current_doc_id = doc_info.get("document_id", "")

                suggs = result.get("suggestions", [])
                # Gắn ui_id riêng để quản lý state nút bấm
                for s in suggs:
                    if "ui_id" not in s:
                        s["ui_id"] = str(uuid.uuid4())

                st.session_state.suggestions = suggs

                if suggs:
                    st.success(f"✅ Tìm thấy **{len(suggs)}** đề xuất liên kết.")
                else:
                    msg = result.get("message", "Không tìm thấy đề xuất liên kết nào.")
                    st.warning(f"⚠️ {msg}")

    # ──────────────────────────────────────────────────────────
    # Nhóm suggestions theo matched_anchor để hiển thị dạng group
    #   - Hiển thị banner tóm tắt 3 cấp ngay sau khi nhận kết quả từ Backend.
    #   - Toggle giữa "Grouped view" (Cấp 1 & 2) và "Individual view" (Cấp 3).
    #   - Grouped view: mỗi anchor_value = 1 container, có nút group-confirm.
    #   - Individual view: mỗi suggestion = 1 card, dùng confirm_link đơn lẻ.
    # ──────────────────────────────────────────────────────────
    if st.session_state.suggestions:
        st.divider()
 
        # ── Bộ lọc phương pháp (giữ nguyên) ─────────────────────────────
        all_methods = sorted({
            s.get("method", "?") for s in st.session_state.suggestions
        })
        selected_methods = st.multiselect(
            "🔎 Lọc theo phương pháp matching",
            options=all_methods,
            default=all_methods,
        )
 
        # ── Banner tóm tắt 3 cấp ─────────────────────────────────────────
        # Hiển thị ngay dưới bộ lọc để user có cái nhìn tổng quan trước khi xét duyệt
        auto_confirmed_list = [
            s for s in st.session_state.suggestions
            if s.get("status") == "AUTO_CONFIRMED"
            and s.get("method") in selected_methods
        ]
        group_pending_list = [
            s for s in st.session_state.suggestions
            if s.get("tier") == 2
            and s.get("method") in selected_methods
            and s.get("ui_id") not in st.session_state.confirmed_uids
            and s.get("ui_id") not in st.session_state.skipped_uids
        ]
        manual_pending_list = [
            s for s in st.session_state.suggestions
            if s.get("tier") == 3
            and s.get("method") in selected_methods
            and s.get("ui_id") not in st.session_state.confirmed_uids
            and s.get("ui_id") not in st.session_state.skipped_uids
        ]
 
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric(
                "⚡ Cấp 1 — Tự động xác nhận",
                len(auto_confirmed_list),
                help="ID_BASED/HYBRID, confidence ≥ 85% — đã ghi vào Neo4j tự động",
            )
        with c2:
            st.metric(
                "🗂️ Cấp 2 — Group-confirm",
                len(group_pending_list),
                help="ID_BASED/HYBRID, confidence 70–85% — bấm 1 lần cho cả nhóm",
            )
        with c3:
            st.metric(
                "🔬 Cấp 3 — Manual review",
                len(manual_pending_list),
                help="SEMANTIC hoặc confidence < 70% — xem xét từng đề xuất",
            )
 
        if auto_confirmed_list:
            st.success(
                f"✅ **{len(auto_confirmed_list)}** liên kết đã được xác nhận tự động "
                f"(ID_BASED/HYBRID, confidence ≥ 85%) và ghi vào Neo4j."
            )

            # Tổng hợp các nhóm đã auto-confirm theo matched_anchor
            auto_summary = {}
            for s in auto_confirmed_list:
                key = s.get("matched_anchor") or s.get("reference") or "Không xác định"
                auto_summary[key] = auto_summary.get(key, 0) + 1

            auto_summary_df = (
                pd.DataFrame(
                    [
                        {"Loại cấu kiện / Reference": k, "Số links": v}
                        for k, v in auto_summary.items()
                    ]
                )
                .sort_values("Số links", ascending=False)
                .reset_index(drop=True)
            )

            with st.expander("📌 Chi tiết các nhóm đã được tự động xác nhận", expanded=True):
                st.markdown(
                    "Đã xét duyệt/tự động xác nhận toàn bộ các nhóm liên kết sau:"
                )
                st.dataframe(auto_summary_df, use_container_width=True)

        # ── Chọn chế độ xem ──────────────────────────────────────────────
        # "Nhóm" phù hợp cho Cấp 2 (ID_BASED/HYBRID số lượng lớn, confidence tốt).
        # "Từng đề xuất" phù hợp cho Cấp 3 (SEMANTIC, cần xem xét kỹ từng cái).
        view_mode = st.radio(
            "Chế độ xem",
            [
                "🗂️ Nhóm theo loại cấu kiện (Cấp 2)",
                "🔬 Từng đề xuất riêng lẻ (Cấp 3 — Semantic)",
            ],
            horizontal=True,
        )
 
        # ── CHẾ ĐỘ NHÓM — Cấp 2 ─────────────────────────────────────────
        if "Nhóm" in view_mode:
            # Lấy suggestions cần group-confirm (tier=2, chưa xử lý)
            # Nếu user muốn xem Cấp 1 trong nhóm để có thể override, bỏ filter tier==2
            candidates = [
                s for s in st.session_state.suggestions
                if s.get("method") in selected_methods
                and s.get("status") != "AUTO_CONFIRMED"   # ẩn auto-confirmed khỏi review
                and s.get("ui_id") not in st.session_state.confirmed_uids
                and s.get("ui_id") not in st.session_state.skipped_uids
            ]
 
            # Nhóm theo matched_anchor (loại cấu kiện)
            grouped: dict = {}
            for s in candidates:
                key = s.get("matched_anchor", "Không xác định")
                grouped.setdefault(key, []).append(s)
 
            total_groups = len(grouped)
            total_pending = len(candidates)
 
            if not grouped:
                st.success(
                    "🎉 Không còn đề xuất Cấp 2 cần xét duyệt thủ công. "
                    "Các liên kết đủ tin cậy đã được tự động xác nhận ở Cấp 1."
                )

                if auto_confirmed_list:
                    auto_summary = {}
                    for s in auto_confirmed_list:
                        key = s.get("matched_anchor") or "Không xác định"
                        auto_summary[key] = auto_summary.get(key, 0) + 1

                    auto_summary_df = (
                        pd.DataFrame(
                            [
                                {"Loại cấu kiện / Reference": k, "Số links": v}
                                for k, v in auto_summary.items()
                            ]
                        )
                        .sort_values("Số links", ascending=False)
                        .reset_index(drop=True)
                    )

                    st.markdown("**Các nhóm đã được tự động xác nhận:**")
                    st.dataframe(auto_summary_df, use_container_width=True)
            else:
                st.subheader(
                    f"💡 {total_groups} nhóm cấu kiện · {total_pending} đề xuất chờ xét duyệt"
                )
 
                # Nút "Đồng ý tất cả" ở cấp toàn bộ (cho user lười 😄)
                col_all1, col_all2, _ = st.columns([2, 2, 4])
                if col_all1.button("✅ Đồng ý TẤT CẢ nhóm", type="primary"):
                    all_ids  = [s["suggestion_id"] for s in candidates if s.get("suggestion_id")]
                    res = api_confirm_group(
                        project_id, all_ids, "CONFIRMED",
                        doc_id=st.session_state.current_doc_id,
                        doc_name=st.session_state.current_file_name,
                        suggestions_data=candidates,
                    )
                    if "error" in res:
                        st.error(f"❌ {res['error']}")
                    else:
                        for s in candidates:
                            st.session_state.confirmed_uids.add(s.get("ui_id", ""))
                        st.toast(f"✅ Đã xác nhận tất cả {total_pending} đề xuất", icon="✅")
                        st.rerun()
 
                if col_all2.button("❌ Bỏ qua TẤT CẢ nhóm"):
                    all_ids = [s["suggestion_id"] for s in candidates if s.get("suggestion_id")]
                    api_confirm_group(
                        project_id, all_ids, "REJECTED",
                        doc_id=st.session_state.current_doc_id,
                        doc_name=st.session_state.current_file_name,
                    )
                    for s in candidates:
                        st.session_state.skipped_uids.add(s.get("ui_id", ""))
                    st.rerun()
 
                st.divider()
 
                # Render từng nhóm — mỗi nhóm = 1 container
                for anchor_value, group in sorted(grouped.items()):
                    methods_in_group = {s.get("method", "?") for s in group}
                    max_conf = max(s.get("confidence", 0) for s in group)
                    min_conf = min(s.get("confidence", 0) for s in group)
 
                    METHOD_BADGE = {"ID_BASED": "🟢", "HYBRID": "🟡", "SEMANTIC": "🔵"}
                    method_display = "  ·  ".join(
                        f"{METHOD_BADGE.get(m, '⚪')} {m}" for m in sorted(methods_in_group)
                    )
 
                    with st.container(border=True):
                        # Header của nhóm
                        col_h, col_btn = st.columns([6, 3])
                        with col_h:
                            st.markdown(
                                f"**🔖 Anchor:** `{anchor_value}`  \n"
                                f"**{len(group)}** cấu kiện  ·  "
                                f"{method_display}  ·  "
                                f"Confidence: `{min_conf:.0%}` – `{max_conf:.0%}`"
                            )
 
                        with col_btn:
                            btn_col1, btn_col2 = st.columns(2)
 
                            # Nút Đồng ý cho cả nhóm này
                            if btn_col1.button(
                                f"✅ Đồng ý ({len(group)})",
                                key=f"grp_yes_{anchor_value}",
                                type="primary",
                                use_container_width=True,
                            ):
                                ids = [s["suggestion_id"] for s in group if s.get("suggestion_id")]
                                with st.spinner(f"Đang lưu {len(group)} liên kết..."):
                                    res = api_confirm_group(
                                        project_id, ids, "CONFIRMED",
                                        doc_id=st.session_state.current_doc_id,
                                        doc_name=st.session_state.current_file_name,
                                        suggestions_data=group,
                                    )
                                if "error" in res:
                                    st.error(f"❌ {res['error']}")
                                else:
                                    for s in group:
                                        st.session_state.confirmed_uids.add(s.get("ui_id", ""))
                                    st.toast(
                                        f"✅ Đã xác nhận {len(group)} cấu kiện `{anchor_value}`",
                                        icon="✅",
                                    )
                                    st.rerun()
 
                            # Nút Bỏ qua cả nhóm này
                            if btn_col2.button(
                                f"❌ Bỏ qua ({len(group)})",
                                key=f"grp_no_{anchor_value}",
                                use_container_width=True,
                            ):
                                ids = [s["suggestion_id"] for s in group if s.get("suggestion_id")]
                                api_confirm_group(
                                    project_id, ids, "REJECTED",
                                    doc_id=st.session_state.current_doc_id,
                                    doc_name=st.session_state.current_file_name,
                                )
                                for s in group:
                                    st.session_state.skipped_uids.add(s.get("ui_id", ""))
                                st.rerun()
 
                        # Chi tiết từng instance — có thể expand để xem hoặc override
                        with st.expander(f"🔍 Xem chi tiết {len(group)} instance"):
                            for item in group:
                                item_conf   = item.get("confidence", 0)
                                item_method = item.get("method", "?")
                                item_gid    = item.get("element_global_id", "?")
                                item_name   = item.get("element_name", "?")
 
                                CONF_BADGE  = "🟢" if item_conf >= 0.80 else ("🟡" if item_conf >= 0.50 else "🔴")
                                st.text(
                                    f"  {CONF_BADGE}  {item_conf:.0%}  |  "
                                    f"{item_method}  |  "
                                    f"{item_name[:45]}  |  "
                                    f"GlobalId: {item_gid}"
                                )
 
        # ── CHẾ ĐỘ TỪNG ĐỀ XUẤT — Cấp 3 (Manual review) ────────────────
        else:
            # Chỉ lấy suggestions cần manual review (tier=3 hoặc tất cả tùy filter)
            manual_candidates = [
                s for s in st.session_state.suggestions
                if s.get("method") in selected_methods
                and s.get("status") != "AUTO_CONFIRMED"
                and s.get("ui_id") not in st.session_state.confirmed_uids
                and s.get("ui_id") not in st.session_state.skipped_uids
            ]
 
            if not manual_candidates:
                st.success("🎉 Đã xét duyệt toàn bộ đề xuất trong phiên này!")
 
            st.subheader(f"🔬 {len(manual_candidates)} đề xuất cần xem xét thủ công")
 
            for s in manual_candidates:
                uid        = s.get("ui_id", str(uuid.uuid4()))
                conf       = s.get("confidence", 0.0)
                method     = s.get("method", "?")
                elem_name  = s.get("element_name", "Unknown")
                elem_gid   = s.get("element_global_id", "")
                anchor_val = s.get("matched_anchor", "")
                evidence   = s.get("evidence", "")
 
                METHOD_BADGE = {"ID_BASED": "🟢", "SEMANTIC": "🔵", "HYBRID": "🟡"}
                CONF_BADGE   = "🟢" if conf >= 0.80 else ("🟡" if conf >= 0.50 else "🔴")
 
                with st.container(border=True):
                    left, right = st.columns([7, 2])
 
                    with left:
                        st.markdown(
                            f"**📄 Tài liệu:** `{st.session_state.current_file_name}`  \n"
                            f"**🧱 Cấu kiện:** `{elem_name}`  \n"
                            f"**🆔 GlobalId:** `{elem_gid}`  \n"
                            f"**🔑 Anchor khớp:** `{anchor_val}`"
                        )
                        st.markdown(
                            f"**Độ tin cậy:** {CONF_BADGE} `{conf:.0%}`  &nbsp;  "
                            f"**Phương pháp:** {METHOD_BADGE.get(method, '⚪')} `{method}`"
                        )
                        if evidence:
                            st.caption(
                                f"📎 **Bằng chứng:** _{evidence[:250]}_"
                                + ("…" if len(evidence) > 250 else "")
                            )
 
                    with right:
                        # Confirm đơn lẻ — dùng confirm_link endpoint cũ
                        if st.button("✅ Đồng ý", key=f"yes_{uid}", type="primary",
                                     use_container_width=True):
                            confirm_payload = {
                                "document_id":       st.session_state.current_doc_id,
                                "document_name":     st.session_state.current_file_name,
                                "element_global_id": elem_gid,
                                "suggestion_id":     s.get("suggestion_id", uid),
                                "confidence":        conf,
                                "method":            method,
                                "evidence":          evidence,
                                "confirmed_by":      "human_via_ui",
                            }
                            with st.spinner("Đang lưu vào Neo4j..."):
                                res = api_confirm_link(project_id, confirm_payload)
                            if "error" in res:
                                st.error(f"❌ {res['error']}")
                            else:
                                st.session_state.confirmed_uids.add(uid)
                                st.toast(f"🎉 Đã lưu: {elem_name}", icon="✅")
                                st.rerun()
 
                        if st.button("❌ Bỏ qua", key=f"no_{uid}",
                                     use_container_width=True):
                            st.session_state.skipped_uids.add(uid)
                            st.rerun()
        st.divider()

# ══════════════════════════════════════════════════════════════
# TAB 3 — Confirmed Links
# ══════════════════════════════════════════════════════════════
with tab_links:
    st.subheader("🔗 Liên kết đã xác nhận trong dự án")
    st.caption(
        "Hiển thị tất cả triple đã được lưu vào Neo4j: "
        "`(Document)-[:REFERENCES]->(Element)`"
    )

    if st.button("🔄 Tải danh sách từ Neo4j", type="primary"):
        with st.spinner("Đang truy vấn..."):
            res = api_get_links(project_id)

        if "error" in res:
            st.error(f"❌ {res['error']}")
        else:
            links = res.get("links", [])
            if not links:
                st.info(f"Chưa có liên kết nào trong project `{project_id}`.")
            else:
                st.success(f"✅ Tìm thấy **{len(links)}** liên kết đã xác nhận.")
                df = pd.DataFrame(links)

                # Đổi tên cột cho đẹp
                col_rename = {
                    "document_name":      "Tài liệu",
                    "element_name":       "Cấu kiện",
                    "element_global_id":  "GlobalId",
                    "element_type":       "IFC Type",
                    "confidence":         "Confidence",
                    "method":             "Phương pháp",
                    "created_at":         "Thời gian",
                    "created_by":         "Bởi",
                }
                df = df.rename(columns=col_rename)

                # Format confidence
                if "Confidence" in df.columns:
                    df["Confidence"] = df["Confidence"].apply(
                        lambda x: f"{float(x):.0%}" if x is not None else "-"
                    )

                display_cols = [c for c in col_rename.values() if c in df.columns]
                st.dataframe(df[display_cols], use_container_width=True)

                # Export CSV
                csv = df.to_csv(index=False).encode("utf-8-sig")
                st.download_button(
                    label="📥 Tải về CSV",
                    data=csv,
                    file_name=f"bim_links_{project_id}.csv",
                    mime="text/csv",
                )

    # Tóm tắt phiên làm việc hiện tại (từ session_state)
    if st.session_state.confirmed_uids:
        st.divider()
        st.subheader("📋 Xác nhận trong phiên này")

        confirmed_this_session = [
            s for s in st.session_state.suggestions
            if s.get("ui_id") in st.session_state.confirmed_uids
        ]
        if confirmed_this_session:
            df_session = pd.DataFrame([
                {
                    "Tài liệu":    st.session_state.current_file_name,
                    "Cấu kiện":    s.get("element_name", "?"),
                    "GlobalId":    s.get("element_global_id", "?"),
                    "Phương pháp": s.get("method", "?"),
                    "Confidence":  f"{s.get('confidence', 0):.0%}",
                }
                for s in confirmed_this_session
            ])
            st.dataframe(df_session, use_container_width=True)

# ══════════════════════════════════════════════════════════════
# TAB 4 — Xuất ICDD / Linkset
# ══════════════════════════════════════════════════════════════
with tab_icdd:

    st.subheader("📦 Xuất ICDD Container theo ISO 21597-1")
    st.caption(
        "Đóng gói toàn bộ liên kết ngữ nghĩa đã xác nhận thành "
        "chuẩn **ISO 21597-1 (ICDD)** để trao đổi liên thông với "
        "các phần mềm BIM khác (Bimspot, Trimble Connect, v.v.)"
    )

    # ── Thông tin nền về ICDD ──────────────────────────────────
    with st.expander(" ICDD là gì? Tại sao cần chuẩn này?", expanded=False):
        st.markdown("""
**ICDD (Information Container for linked Document Delivery)** là chuẩn ISO 21597-1:2020,
định nghĩa cách đóng gói thông tin BIM để trao đổi giữa các bên trong dự án.

Một ICDD package là một file ZIP (đổi đuôi thành `.icdd`) với 4 thư mục cố định:

| Thư mục | Nội dung | Bắt buộc? |
|---|---|---|
| `Container.rdf` | File index mô tả toàn bộ package | ✅ Bắt buộc |
| `Payload Documents/` | Tài liệu gốc: PDF, DOCX, XLSX, IFC | ✅ Bắt buộc |
| `Linkset Documents/` | File `.ttl` chứa liên kết ngữ nghĩa | ✅ Bắt buộc |
| `Ontology Resources/` | Ontology tham chiếu: bot.ttl, IFC2X3_Final.ttl | 📌 Khuyến nghị |

Namespace của ứng dụng này dùng domain thật:
`https://bim-semantic-linker.onrender.com/` — đây là **Linked Data URI** hợp lệ, 
không phải URN tạm thời.
        """)

    st.divider()

    # ══════════════════════════════════════════════════════════
    # PHẦN A — Export Linkset (file RDF mô tả liên kết)
    # ══════════════════════════════════════════════════════════
    st.markdown("### 📄 Phần A — Export Linkset")
    st.caption(
        "Linkset là file RDF liệt kê tất cả liên kết "
        "`(Document)–[:REFERENCES]→(Element)` đã được xác nhận, "
        "tuân thủ cấu trúc `ls:Linkset → ls:1-to-1Link → ls:LinkElement` "
        "của ISO 21597-1."
    )

    col_fmt, col_btn_ls = st.columns([2, 3])

    with col_fmt:
        # Người dùng chọn format serialize
        fmt_choice = st.selectbox(
            "Định dạng serialize",
            options=["turtle", "json-ld", "xml"],
            format_func=lambda x: {
                "turtle":  "🐢 Turtle (.ttl) — Dễ đọc, khuyến nghị",
                "json-ld": "📋 JSON-LD (.jsonld) — Tốt cho web API",
                "xml":     "📰 RDF/XML (.rdf) — Tương thích rộng nhất",
            }[x],
            help=(
                "Turtle: dễ đọc bằng mắt thường, phù hợp lưu file vật lý.\n"
                "JSON-LD: dễ xử lý bằng Python/JavaScript, phù hợp cho API.\n"
                "RDF/XML: tương thích rộng nhất với các công cụ cũ."
            ),
        )

    ext_map   = {"turtle": "ttl", "json-ld": "jsonld", "xml": "rdf"}
    mime_map  = {
        "turtle":  "text/turtle",
        "json-ld": "application/ld+json",
        "xml":     "application/rdf+xml",
    }

    with col_btn_ls:
        st.write("")  # Căn chỉnh vertical
        st.write("")
        if st.button("🔄 Tạo Linkset từ Neo4j", type="primary",
                     use_container_width=True, key="gen_linkset"):
            with st.spinner("Đang query Neo4j và serialize RDF..."):
                resp = api_export_linkset(project_id, fmt=fmt_choice)
            if resp is not None:
                # Lưu vào session_state để hiển thị preview và nút download
                st.session_state["linkset_bytes"]    = resp.content
                st.session_state["linkset_fmt"]      = fmt_choice
                st.session_state["linkset_filename"] = (
                    f"linkset_{project_id}.{ext_map[fmt_choice]}"
                )
                st.success(
                    f"✅ Linkset được tạo thành công! "
                    f"Kích thước: **{len(resp.content) / 1024:.1f} KB**"
                )

    # Preview và nút download — chỉ hiện sau khi đã generate
    if "linkset_bytes" in st.session_state and st.session_state["linkset_bytes"]:
        linkset_str = st.session_state["linkset_bytes"].decode("utf-8")

        with st.expander(
            f"👁️ Xem trước nội dung "
            f"`{st.session_state['linkset_filename']}` (200 dòng đầu)",
            expanded=False
        ):
            # Hiển thị với syntax highlighting giả lập qua code block
            preview_lines = linkset_str.split("\n")[:200]
            lang = "turtle" if st.session_state["linkset_fmt"] == "turtle" else "json"
            st.code("\n".join(preview_lines), language=lang)
            if len(linkset_str.split("\n")) > 200:
                st.caption(
                    f"_(Hiển thị 200/{len(linkset_str.split(chr(10)))} dòng. "
                    "Tải file để xem đầy đủ.)_"
                )

        st.download_button(
            label=f"📥 Tải xuống {st.session_state['linkset_filename']}",
            data=st.session_state["linkset_bytes"],
            file_name=st.session_state["linkset_filename"],
            mime=mime_map[st.session_state["linkset_fmt"]],
            type="primary",
            use_container_width=True,
        )

    st.divider()

    # ══════════════════════════════════════════════════════════
    # PHẦN B — Export Container.rdf (file index bắt buộc)
    # ══════════════════════════════════════════════════════════
    st.markdown("### 🗂️ Phần B — Tạo Container.rdf")
    st.caption(
        "Container.rdf là file INDEX bắt buộc nằm ở root của ICDD package. "
        "Nó liệt kê tất cả tài liệu, linkset, và ontology có trong package, "
        "tương tự như file `manifest.json` trong các định dạng đóng gói khác."
    )

    if st.button("🔄 Tạo Container.rdf", use_container_width=True,
                 key="gen_container"):
        with st.spinner("Đang lấy danh sách tài liệu từ Neo4j..."):
            resp = api_export_container(project_id)
        if resp is not None:
            st.session_state["container_bytes"] = resp.content
            st.success(
                f"✅ Container.rdf được tạo thành công! "
                f"Kích thước: **{len(resp.content) / 1024:.1f} KB**"
            )

    if "container_bytes" in st.session_state and st.session_state["container_bytes"]:
        with st.expander("👁️ Xem trước nội dung Container.rdf", expanded=False):
            container_str = st.session_state["container_bytes"].decode("utf-8")
            st.code(container_str[:5000], language="xml")
            if len(container_str) > 5000:
                st.caption("_(Chỉ hiển thị 5000 ký tự đầu.)_")

        st.download_button(
            label="📥 Tải xuống Container.rdf",
            data=st.session_state["container_bytes"],
            file_name="Container.rdf",
            mime="application/rdf+xml",
            type="primary",
            use_container_width=True,
        )

    st.divider()

    # ══════════════════════════════════════════════════════════
    # PHẦN C — Workflow 4 bước: Starter ZIP → bổ sung file → validate → xuất .icdd
    # ══════════════════════════════════════════════════════════
    st.markdown("### 📦 Phần C — Đóng gói ICDD hoàn chỉnh")

    with st.expander("📋 Quy trình 4 bước để có ICDD package hoàn chỉnh 100%", expanded=True):
        st.markdown(f"""
Workflow tránh thao tác thủ công, chỉ làm việc với file ZIP trong bước bổ sung nội dung; app sẽ kiểm tra và xuất `.icdd` cuối cùng.

| Bước | Việc cần làm | Kết quả |
|---|---|---|
| **1** | Tạo và tải bộ khung metadata-only dạng ZIP | `SNT_ICDD_starter.zip` hoặc `{project_id}_ICDD_starter.zip` |
| **2** | Giải nén ZIP, bổ sung tài liệu gốc và ontology | `Payload Documents/` và `Ontology Resources/` không còn rỗng |
| **3** | Upload lại ZIP đã bổ sung vào app | App kiểm tra cấu trúc ICDD |
| **4** | Nếu hợp lệ, tải file `.icdd` cuối cùng | `{project_id}_ICDD_package.icdd` |

Bộ khung ZIP ở bước 1 chứa sẵn:
- ✅ `Container.rdf` — metadata/index tự động từ Neo4j
- ✅ `Linkset Documents/linkset_{project_id}.ttl` — linkset tự động từ Neo4j
- ⚠️ `Payload Documents/` — cần thêm tài liệu gốc: PDF, DOCX, XLSX, IFC
- ⚠️ `Ontology Resources/` — cần thêm ontology: `bot.ttl`, `IFC2X3_Final.ttl`, v.v.
        """)

    st.markdown("#### Bước 1 — Tải bộ khung ICDD metadata-only package dạng ZIP")
    st.caption(
        "File tải xuống là ZIP để bạn giải nén và bổ sung nội dung dễ dàng trên Windows. "
        "Đây chưa phải package ICDD hoàn chỉnh vì chưa chứa tài liệu gốc và ontology."
    )

    col_starter, col_starter_info = st.columns([1, 1])

    with col_starter:
        if st.button(
            "🚀 Tạo bộ khung ICDD metadata-only package",
            type="primary",
            use_container_width=True,
            key="gen_icdd_starter_zip",
        ):
            with st.spinner(
                "Đang tạo bộ khung ICDD metadata-only... "
                "(query Neo4j → tạo Container.rdf + Linkset → đóng gói ZIP)"
            ):
                resp = api_package_icdd(project_id)

            if resp is not None:
                st.session_state["icdd_starter_zip_bytes"] = resp.content
                st.success(
                    f"✅ Bộ khung ICDD metadata-only đã sẵn sàng! "
                    f"Kích thước ZIP: **{len(resp.content) / 1024:.1f} KB**"
                )

    with col_starter_info:
        st.info(
            "💡 tải `.zip`, bổ sung file, rồi upload lại để app tự xuất `.icdd`."
        )

    if "icdd_starter_zip_bytes" in st.session_state and st.session_state["icdd_starter_zip_bytes"]:
        st.download_button(
            label=f"📥 Tải bộ khung ICDD metadata-only package dạng ZIP — {project_id}_ICDD_starter.zip",
            data=st.session_state["icdd_starter_zip_bytes"],
            file_name=f"{project_id}_ICDD_starter.zip",
            mime="application/zip",
            type="primary",
            use_container_width=True,
        )

    st.markdown("#### Bước 2 — Giải nén ZIP và bổ sung file còn thiếu")
    st.markdown("""
Sau khi tải `*_ICDD_starter.zip`, hãy giải nén và bổ sung:

**Thêm vào `Payload Documents/`:**
```text
Biên bản Nghiệm thu Vật liệu.pdf
SNT-DEF-ARC-DOOR-L1-MAIN-v1.pdf
Nhật_ký_thi_công_20_05_2026.docx
Snowdon Towers Sample Structural.ifc
... các tài liệu gốc khác của dự án
```

**Thêm vào `Ontology Resources/`:**
```text
bot.ttl
IFC2X3_Final.ttl
... ontology tham chiếu khác nếu có
```

Sau đó nén lại thành một file `.zip`.
    """)

    st.markdown("#### Bước 3 — Upload ZIP đã bổ sung Payload Documents và Ontology Resources")
    completed_zip = st.file_uploader(
        "Tải lên ZIP đã bổ sung file",
        type=["zip"],
        key="completed_icdd_zip_upload",
        help=(
            "ZIP cần có Container.rdf ở root, Linkset Documents/, Payload Documents/, "
            "và Ontology Resources/ để app có thể tự chuẩn hoá."
        ),
    )

    if completed_zip is not None:
        zip_bytes = completed_zip.getvalue()
        with st.spinner("Đang kiểm tra cấu trúc ICDD trong ZIP..."):
            validation = validate_and_normalize_icdd_zip(zip_bytes, project_id)

        summary = validation.get("summary", {})
        if summary:
            st.write("**Tóm tắt nội dung phát hiện trong ZIP:**")
            st.dataframe(
                pd.DataFrame(
                    [
                        {"Thành phần": k, "Số lượng": v}
                        for k, v in summary.items()
                    ]
                ),
                use_container_width=True,
                hide_index=True,
            )

        for warning in validation.get("warnings", []):
            st.warning(f"⚠️ {warning}")

        if not validation.get("ok"):
            st.error("❌ ZIP chưa đủ điều kiện để xuất ICDD package hoàn chỉnh.")
            for err in validation.get("errors", []):
                st.markdown(f"- {err}")
        else:
            st.session_state["validated_icdd_bytes"] = validation["icdd_bytes"]
            st.success("✅ Package đã đủ cấu trúc ICDD")

            with st.expander("👁️ Xem danh sách file sẽ được đóng gói vào `.icdd`", expanded=False):
                files_df = pd.DataFrame(
                    [{"Đường dẫn trong package": n} for n in validation.get("files", [])]
                )
                st.dataframe(files_df, use_container_width=True, hide_index=True)

    st.markdown("#### Bước 4 — Tải xuống ICDD package hoàn chỉnh")
    if "validated_icdd_bytes" in st.session_state and st.session_state["validated_icdd_bytes"]:
        st.success("✅ Package đã đủ cấu trúc ICDD")
        st.download_button(
            label=f"📥 Tải xuống {project_id}_ICDD_package.icdd",
            data=st.session_state["validated_icdd_bytes"],
            file_name=f"{project_id}_ICDD_package.icdd",
            mime="application/zip",
            type="primary",
            use_container_width=True,
        )
    else:
        st.info(
            "Đã upload ZIP hợp lệ ở Bước 3, nút tải `.icdd` sẽ xuất hiện."
        )

    st.divider()

    # ══════════════════════════════════════════════════════════
    # PHẦN D — Roadmap: Những gì CHƯA đủ và bước tiếp theo
    # ══════════════════════════════════════════════════════════
    st.markdown("### 🗺️ Roadmap: Từ Linkset đến trải nghiệm tương tác đầy đủ")

    st.markdown("""
File Linkset `.ttl` bạn vừa tạo là **nền tảng ngữ nghĩa** — nhưng ba use case 
tương tác nâng cao cần thêm các lớp hạ tầng bổ sung. 
Bảng dưới đây cho thấy mức độ sẵn sàng hiện tại:
    """)

    roadmap_data = {
        "Use case": [
            "Click GlobalId → xem tất cả tài liệu liên quan",
            "Tìm kiếm: tài liệu nào nhắc đến cấu kiện X?",
            "Highlight text trong PDF khi đề cập đến cấu kiện",
            "Click text → 3D viewer zoom đến cấu kiện",
        ],
        "Trạng thái": [
            "✅ Sẵn sàng",
            "✅ Sẵn sàng",
            "🔶 Cần thêm dữ liệu",
            "🔴 Cần xây thêm lớp mới",
        ],
        "Cần thêm gì": [
            "Thêm endpoint GET /elements/<globalId>/documents",
            "Query Neo4j MATCH (d)-[:REFERENCES]->(e {globalId: $gid})",
            "Lưu thêm page + text_bbox vào field 'evidence' trong Neo4j",
            "Geometry từ IFC + tích hợp IFC.js hoặc xeokit viewer",
        ],
        "Độ khó": ["Dễ (1 ngày)", "Dễ (1 ngày)", "Trung bình (1 tuần)", "Khó (3-4 tuần)"],
    }

    import pandas as pd
    df_roadmap = pd.DataFrame(roadmap_data)
    st.dataframe(
        df_roadmap,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Use case":      st.column_config.TextColumn(width="large"),
            "Trạng thái":    st.column_config.TextColumn(width="medium"),
            "Cần thêm gì":   st.column_config.TextColumn(width="large"),
            "Độ khó":        st.column_config.TextColumn(width="medium"),
        }
    )

    st.info(
        "**Bước ngắn hạn được đề xuất:** Thêm endpoint "
        "`GET /api/projects/<pid>/elements/<globalId>/documents` "
        "để enable Use case 1 & 2 ngay lập tức — đây là tính năng "
        "có giá trị thực tế cao nhất và chi phí triển khai thấp nhất."
    )
